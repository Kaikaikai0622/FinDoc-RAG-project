"""DOCX 解析模块

使用 python-docx 解析 .docx 文件。
- Heading 样式段落 → category="Title"
- 普通段落 → category="Text"
- 追踪最近一个 Heading 生成 section_title 注入到 metadata
- page_or_index 使用段落顺序索引（python-docx 无法获取真实页码）

NOTE (v2 TODO): doc.paragraphs 不包含表格内容，
                后续版本需遍历 doc.tables 生成 category="Table" 元素。
"""
import logging
from pathlib import Path
from typing import List

from .models import ParsedDocument, ParsedElement

logger = logging.getLogger(__name__)

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    docx = None  # type: ignore[assignment]
    DOCX_AVAILABLE = False
    logger.warning("python-docx 库不可用，无法解析 .docx 文件。请执行: pip install python-docx")


class DocxParser:
    """DOCX 解析器"""

    def parse(self, file_path: str) -> ParsedDocument:
        """解析 .docx 文件

        Args:
            file_path: .docx 文件路径

        Returns:
            ParsedDocument 对象

        Raises:
            RuntimeError: python-docx 不可用
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError(
                "python-docx 库不可用，无法解析 .docx 文件。请执行: pip install python-docx"
            )

        path = Path(file_path)
        source_file = path.name
        elements = self._parse_with_python_docx(file_path)

        return ParsedDocument(
            source_file=source_file,
            file_type="docx",
            elements=elements,
        )

    # ------------------------------------------------------------------
    # 内部方法
    # ------------------------------------------------------------------

    def _parse_with_python_docx(self, file_path: str) -> List[ParsedElement]:
        """遍历段落生成 ParsedElement 列表"""
        elements: List[ParsedElement] = []
        document = docx.Document(file_path)

        current_section_title = ""

        for para_idx, para in enumerate(document.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            is_heading = style_name.startswith("Heading") or style_name.startswith("标题")

            if is_heading:
                current_section_title = text
                elements.append(ParsedElement(
                    text=text,
                    category="Title",
                    page_or_index=para_idx,
                    metadata={
                        "paragraph_index": para_idx,
                        "style": style_name,
                        "section_title": text,
                    },
                ))
            else:
                elements.append(ParsedElement(
                    text=text,
                    category="Text",
                    page_or_index=para_idx,
                    metadata={
                        "paragraph_index": para_idx,
                        "style": style_name,
                        "section_title": current_section_title,
                    },
                ))

        # TODO v2: 遍历 document.tables 生成 category="Table" 元素
        # for table in document.tables:
        #     ...

        return elements


def parse_docx(file_path: str) -> ParsedDocument:
    """解析 DOCX 文件的便捷函数"""
    parser = DocxParser()
    return parser.parse(file_path)


